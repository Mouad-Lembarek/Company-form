from flask import Flask, render_template, redirect, url_for, session
from flask_wtf import FlaskForm
from wtforms import *
from wtforms.validators import *

from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = '123456'

class WordForm(FlaskForm):
    companyName = StringField([DataRequired()])
    companyOwner = StringField([DataRequired()])
    companyAddress = StringField([DataRequired()])
    companyCreationDate = StringField([DataRequired()])
    submit = SubmitField('Generate Document')


@app.route('/', methods=['POST', 'GET'])
def index():
    form = WordForm()
    if form.validate_on_submit():
        # COLLECTION DES INFORMATIONS ENTREES
        companyName = form.companyName.data
        companyOwner = form.companyOwner.data
        companyAddress = form.companyAddress.data
        companyCreationDate = form.companyCreationDate.data

        # DOCUMENTATION (https://python-docx.readthedocs.io/en/latest/)
        # Creating Word Document
        document = Document()

        # L'AJOUT DU TITRE 
        document.add_heading(companyName, 0)

        # PREPARATION DES INFORMATION
        data = (
        ('Company Name', companyName),
        ('Company Owner', companyOwner),
        ('Company Address', companyAddress),
        ('Company Creation Date', companyCreationDate)
        )

        # L'AJOUT D'UN TABLEAU
        table = document.add_table(rows=1, cols=2)
        
        # REMPLISSAGE DES INFORMATIONS DANS LE TABLEAU
        for info, infoData in data:
            row = table.add_row().cells
            row[0].text = info
            row[1].text = infoData

        document.save(f'./Docs/{companyName}.docx')
        

    return render_template('form.html', form=form)

app.run(debug=True)